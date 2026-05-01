"""
resume_parser.py
----------------
Handles text extraction from uploaded resume files (PDF and DOCX).

Responsibilities:
  - Detect file type from extension and MIME type
  - Extract raw text while preserving logical line/page structure
  - Provide pdfminer fallback when pdfplumber yields empty pages
  - Raise typed, descriptive exceptions for the API layer to handle

Does NOT perform any cleaning or skill extraction — that is the
responsibility of text_cleaner.py and skill_extractor.py respectively.
"""

import io
import logging
import os
from enum import Enum
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.pdfparser import PDFSyntaxError
from dataclasses import dataclass

logger = logging.getLogger(__name__)

@dataclass
class ParsedResume:
    raw_text: str
    pages: list[str]
    file_type: str
    page_count: int

    @property
    def is_empty(self) -> bool:
        return not self.raw_text.strip()


# ---------------------------------------------------------------------------
# Supported file types
# ---------------------------------------------------------------------------

class FileType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"


SUPPORTED_EXTENSIONS: dict[str, FileType] = {
    ".pdf": FileType.PDF,
    ".docx": FileType.DOCX,
}

MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------

class ResumeParserError(Exception):
    """Base class for all parser errors."""


class UnsupportedFileTypeError(ResumeParserError):
    """Raised when the uploaded file extension is not supported."""


class FileTooLargeError(ResumeParserError):
    """Raised when the uploaded file exceeds the size limit."""


class EmptyDocumentError(ResumeParserError):
    """Raised when no extractable text is found in the document."""


class CorruptedFileError(ResumeParserError):
    """Raised when the file cannot be parsed (corrupted or password-protected)."""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_resume(file_path: str | Path) -> str:
    """
    Extract plain text from a PDF or DOCX resume.

    Args:
        file_path: Absolute or relative path to the uploaded resume file.

    Returns:
        Raw extracted text as a single string. Lines are separated by '\\n'.
        Pages (for PDFs) are separated by '\\n\\n'.

    Raises:
        UnsupportedFileTypeError: File extension is not .pdf or .docx.
        FileTooLargeError: File exceeds MAX_FILE_SIZE_BYTES.
        EmptyDocumentError: No text could be extracted.
        CorruptedFileError: File is unreadable or password-protected.
        ResumeParserError: Any other parsing failure.
    """
    path = Path(file_path)

    _validate_file(path)

    file_type = SUPPORTED_EXTENSIONS[path.suffix.lower()]
    logger.info("Parsing %s file: %s", file_type.value.upper(), path.name)

    if file_type == FileType.PDF:
        raw_text = _extract_pdf(path)
    else:
        raw_text = _extract_docx(path)

    if not raw_text or not raw_text.strip():
        raise EmptyDocumentError(
            f"No extractable text found in '{path.name}'. "
            "The file may be scanned/image-based or contain only graphics."
        )

    logger.info("Extracted %d characters from '%s'", len(raw_text), path.name)
    return raw_text


def parse_resume_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract text directly from an in-memory file (e.g. from FastAPI UploadFile).

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename (used to determine file type).

    Returns:
        Raw extracted text string.

    Raises:
        Same exceptions as parse_resume().
    """
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"File '{filename}' is {len(file_bytes) / 1024 / 1024:.1f} MB. "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB} MB."
        )

    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File type '{suffix}' is not supported. "
            f"Accepted formats: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    file_type = SUPPORTED_EXTENSIONS[suffix]
    buffer = io.BytesIO(file_bytes)

    logger.info("Parsing %s from bytes: %s", file_type.value.upper(), filename)

    if file_type == FileType.PDF:
        raw_text = _extract_pdf_from_buffer(buffer)
    else:
        raw_text = _extract_docx_from_buffer(buffer)

    if not raw_text or not raw_text.strip():
        raise EmptyDocumentError(
            f"No extractable text found in '{filename}'. "
            "The file may be scanned/image-based or contain only graphics."
        )

    logger.info("Extracted %d characters from '%s'", len(raw_text), filename)
    return raw_text


def get_file_type(filename: str) -> FileType:
    """Return the FileType enum for a given filename."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(f"Unsupported file type: '{suffix}'")
    return SUPPORTED_EXTENSIONS[suffix]


# ---------------------------------------------------------------------------
# Internal helpers — validation
# ---------------------------------------------------------------------------

def _validate_file(path: Path) -> None:
    """Run pre-extraction checks on the file."""
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: '{path}'")

    if not path.is_file():
        raise ResumeParserError(f"Path is not a file: '{path}'")

    file_size = os.path.getsize(path)
    if file_size > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"File '{path.name}' is {file_size / 1024 / 1024:.1f} MB. "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB} MB."
        )

    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"File type '{path.suffix}' is not supported. "
            f"Accepted formats: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )


# ---------------------------------------------------------------------------
# Internal helpers — PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file on disk."""
    with open(path, "rb") as f:
        return _extract_pdf_from_buffer(io.BytesIO(f.read()))


def _fitz_extract(buffer: io.BytesIO) -> str:
    """
    Extremely fast PDF text extractor using PyMuPDF (fitz).
    Takes <0.1 seconds even for complex vector layouts.
    """
    buffer.seek(0)
    try:
        import fitz
        import time
        t0 = time.perf_counter()
        # Must read the bytes explicitly for stream=...
        raw_bytes = buffer.read()
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        text_pages = [page.get_text() for page in doc]
        doc.close()
        result = "\n\n".join(text_pages)
        elapsed = (time.perf_counter() - t0) * 1000
        logger.info("PyMuPDF extracted %d chars in %.0f ms", len(result.strip()), elapsed)
        return result
    except ImportError:
        logger.warning("PyMuPDF (fitz) is not installed. Falling back to other parsers.")
        return ""
    except Exception as e:
        logger.warning("PyMuPDF extraction failed: %s", e)
        return ""


def _extract_pdf_from_buffer(buffer: io.BytesIO) -> str:
    """
    Extract text from a PDF buffer using PyMuPDF (fitz) as primary extractor.
    Falls back to pdfminer, then pdfplumber if necessary.
    """
    text = _fitz_extract(buffer)
    if len(text.strip()) >= 50:
        return text

    logger.info("PyMuPDF extracted < 50 chars; trying pdfminer fallback")
    buffer.seek(0)
    text = _pdfminer_extract(buffer)

    # If pdfminer returned very little content, try pdfplumber
    if len(text.strip()) < 50:
        logger.info("pdfminer extracted < 50 chars; trying pdfplumber fallback")
        buffer.seek(0)
        text_fallback = _pdfplumber_extract(buffer)
        if len(text_fallback.strip()) > len(text.strip()):
            logger.info("pdfplumber fallback produced better results")
            return text_fallback

    return text


def _pdfplumber_extract(buffer: io.BytesIO) -> str:
    """
    Primary PDF extractor using pdfplumber.
    Extracts text page-by-page; pages separated by double newline.
    """
    buffer.seek(0)
    pages: list[str] = []

    try:
        with pdfplumber.open(buffer) as pdf:
            if not pdf.pages:
                raise EmptyDocumentError("PDF contains no pages.")

            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text(
                        x_tolerance=3,
                        y_tolerance=3,
                    )
                    if page_text:
                        pages.append(page_text.strip())
                except Exception as e:
                    logger.warning("Failed to extract page %d: %s", i + 1, e)
                    continue

    except PDFSyntaxError as e:
        raise CorruptedFileError(
            f"PDF file appears to be corrupted or password-protected: {e}"
        ) from e
    except Exception as e:
        raise ResumeParserError(f"Unexpected error reading PDF: {e}") from e

    return "\n\n".join(pages)


def _pdfminer_extract(buffer: io.BytesIO) -> str:
    """
    Fallback PDF extractor using pdfminer.
    More robust for unusual fonts and non-standard PDFs.
    """
    buffer.seek(0)
    try:
        text = pdfminer_extract_text(
            buffer,
            page_numbers=None,   # Extract all pages
            maxpages=0,          # No page limit
            caching=True,
            codec="utf-8",
            laparams=None,       # Auto-detect layout parameters
        )
        return text or ""
    except Exception as e:
        logger.warning("pdfminer extraction failed: %s", e)
        return ""


# ---------------------------------------------------------------------------
# Internal helpers — DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file on disk."""
    with open(path, "rb") as f:
        return _extract_docx_from_buffer(io.BytesIO(f.read()))


def _extract_docx_from_buffer(buffer: io.BytesIO) -> str:
    """
    Extract text from a DOCX buffer.

    Extracts from three sources (in order of importance):
      1. Main body paragraphs — primary content and headings
      2. Tables — many resumes put skills in table cells
      3. Text boxes (inline shapes) — some templates use floating text boxes

    Preserves paragraph breaks and heading structure so that the skill
    extractor can reliably detect section headers.
    """
    buffer.seek(0)
    lines: list[str] = []

    try:
        doc = Document(buffer)
    except Exception as e:
        raise CorruptedFileError(
            f"DOCX file appears to be corrupted or invalid: {e}"
        ) from e

    # 1. Main body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 2. Tables (each cell becomes a line)
    for table in doc.tables:
        for row in table.rows:
            row_cells: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                # Join cells with separator so comma-split still works later
                lines.append(" | ".join(row_cells))

    # 3. Text boxes embedded in shapes (common in designed resume templates)
    for shape in _iter_docx_shapes(doc):
        if shape.strip():
            lines.append(shape.strip())

    if not lines:
        raise EmptyDocumentError(
            "DOCX file contains no readable text. "
            "It may contain only images or use an unsupported format."
        )

    return "\n".join(lines)


def _iter_docx_shapes(doc: Document):
    """
    Yield text from drawing objects (text boxes) embedded in the DOCX body.
    These are stored in w:drawing elements and are often missed by basic parsers.
    """
    try:
        from lxml import etree  # python-docx depends on lxml

        # Namespace map for DOCX XML
        nsmap = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        }

        body = doc.element.body
        # Find all text runs inside drawing elements
        for drawing in body.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            texts = drawing.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}t"
            )
            for t in texts:
                if t.text:
                    yield t.text
    except Exception as e:
        # Non-critical — log and continue without shape text
        logger.debug("Could not extract shape text from DOCX: %s", e)